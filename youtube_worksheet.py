from youtube_transcript_api import YouTubeTranscriptApi
import google.generativeai as genai
from googleapiclient.discovery import build
from docx import Document
import re
import os

class YouTubeWorksheet:
    def __init__(self, gemini_api_key, youtube_api_key):
        # Gemini API 초기화
        genai.configure(api_key=gemini_api_key)
        self.model = genai.GenerativeModel('gemini-1.5-pro')
        
        # YouTube API 초기화
        self.youtube = build('youtube', 'v3', developerKey=youtube_api_key)
        
    def get_video_id(self, url):
        # YouTube URL에서 video ID 추출
        video_id = re.search(r'(?:v=|\/)([0-9A-Za-z_-]{11}).*', url)
        return video_id.group(1) if video_id else None

    def get_transcript(self, url):
        video_id = self.get_video_id(url)
        if not video_id:
            return None
            
        try:
            # YouTube API로 비디오 정보 확인
            video_response = self.youtube.videos().list(
                part='snippet',
                id=video_id
            ).execute()
            
            if not video_response['items']:
                return None
                
            # 자막 가져오기
            transcript_list = YouTubeTranscriptApi.list_transcripts(video_id)
            transcript = transcript_list.find_transcript(['en', 'ko'])
            return ' '.join([entry['text'] for entry in transcript.fetch()])
            
        except Exception as e:
            print(f"자막 추출 오류: {e}")
            print(f"비디오 ID: {video_id}")
            return None

    def create_worksheet(self, transcript):
        prompt = f"""
        다음 텍스트의 각 문장에 대해 학습 워크시트를 만들어주세요:
        
        1. 영어 문장에서 중요한 단어나 구를 ___로 대체하여 빈칸 문제를 만드세요
        2. 각 문장의 한국어 번역을 제공하세요
        3. 빈칸에 들어갈 정답(단어나 구)을 별도로 표시하세요
        
        형식:
            - 첫 번째 열: 빈칸이 있는 영어 문장
            - 두 번째 열: 한국어 번역
            - 세 번째 열: 빈칸에 들어갈 정답 단어나 구
        
        예시:
            It is ___ to meet you.|당신을 만나서 반갑습니다.|nice
            I ___ to school every day.|나는 매일 학교에 갑니다.|go
        
        텍스트: {transcript}
        
        표 형식으로 출력:
        문제|한국어 번역|정답
        """
        
        response = self.model.generate_content(prompt)
        return response.text

    def save_to_docx(self, content, output_file="worksheet.docx"):
        doc = Document()
        doc.add_heading('YouTube 학습 활동지', 0)
        
        # 표 생성 및 내용 추가
        rows = content.strip().split('\n')[2:]  # 헤더 제외
        table = doc.add_table(rows=len(rows)+1, cols=3)
        table.style = 'Table Grid'
        
        # 헤더 추가
        headers = ['빈제', '한국어 번역', '정답']
        for i, header in enumerate(headers):
            table.cell(0, i).text = header
            
        # 내용 추가
        for i, row in enumerate(rows):
            cells = row.split('|')
            for j, cell in enumerate(cells):
                table.cell(i+1, j).text = cell.strip()
                
        doc.save(output_file)
        return output_file 
